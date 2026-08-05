"""
Insert English translations into a .docx file, handling merged cells correctly.
Uses XPath-based deduplication to prevent duplicate insertions in merged cells.
Skips cells where English text already exists above the Chinese (avoids redundant
duplicates when the original document already has bilingual EN/CH labels).
"""
import json
import sys
from lxml import etree
import docx
from docx.oxml.ns import qn


def has_chinese(text):
    """Check if text contains Chinese characters."""
    if not text:
        return False
    return any(ord(c) > 0x4e00 for c in text)


def has_english(text):
    """Check if text contains ASCII letters."""
    if not text:
        return False
    return any(c.isalpha() and ord(c) < 128 for c in text)


def is_already_bilingual_v2(text):
    """
    Check if text already contains English translation alongside Chinese.

    Returns True if the text is already bilingual (should skip translation).
    Returns False if the text needs translation.

    Uses multi-layer heuristics to distinguish genuinely bilingual text from
    Chinese text that merely contains short English abbreviations:

    1. Newline-separated bilingual:
       If text has newlines and contains an English-only line AND a Chinese
       line, it's already bilingual (EN translation on its own line).
       e.g. "Equipment Qualification\n设备确认"

    2. For single-line text, uses Chinese character ratio:
       - CN ratio > 0.65 → predominantly Chinese → needs translation
         e.g. "负责向QA申请" (QA is just an abbreviation)
       - CN ratio < 0.40 → predominantly English → already bilingual
         e.g. "Analytical Development (方法开发, AD)"
       - CN ratio 0.40-0.65 → ambiguous zone:
         * Short text (<= 40 chars) → likely bilingual label → skip
           e.g. "Doc. Title 文件标题"
         * Long text (> 40 chars) → likely needs translation
           e.g. "使用范围适用于公司内部所有GMP文件的管理..."
    """
    if not text or not has_chinese(text):
        return False

    # Layer 1: Newline-separated bilingual (most reliable signal)
    if '\n' in text:
        lines = text.split('\n')
        has_en_line = any(
            has_english(line) and not has_chinese(line)
            for line in lines
        )
        has_cn_line = any(has_chinese(line) for line in lines)
        if has_en_line and has_cn_line:
            return True  # Already bilingual with EN/CN on separate lines

    # Layer 2: CN character ratio for single-line text
    cn_chars = sum(1 for c in text if ord(c) > 0x4e00)
    en_chars = sum(1 for c in text if c.isalpha() and ord(c) < 128)
    meaningful = cn_chars + en_chars

    if meaningful == 0:
        return has_chinese(text)

    cn_ratio = cn_chars / meaningful

    # Predominantly Chinese → needs translation
    if cn_ratio > 0.65:
        return False

    # Predominantly English → already bilingual
    if cn_ratio < 0.40:
        return True

    # Ambiguous zone: use text length to decide
    # Short texts are likely bilingual labels; long texts need translation
    if len(text) > 40:
        return False
    else:
        return True


def _safe_para_text(para):
    """Get paragraph text safely, handling runs with None text."""
    try:
        return para.text.strip()
    except (TypeError, AttributeError):
        parts = []
        for elem in para._element:
            if elem.tag.endswith("}t"):
                if elem.text:
                    parts.append(elem.text)
            elif elem.tag.endswith("}r"):
                for sub in elem:
                    if sub.tag.endswith("}t") and sub.text:
                        parts.append(sub.text)
        return "".join(parts).strip()


def cell_has_existing_english_above(cell, target_pi):
    """
    Check if a table cell already has an English-only paragraph above the
    target paragraph index. If so, the cell is already bilingual (EN above,
    CH below) and we should not add another English translation.

    This handles the common pattern in bilingual SOP documents where table
    headers already have English on line 1 and Chinese on line 2.
    """
    if target_pi == 0:
        return False
    for pi in range(target_pi):
        text = _safe_para_text(cell.paragraphs[pi]) if pi < len(cell.paragraphs) else ""
        if text and not has_chinese(text) and has_english(text):
            return True
    return False


def translation_already_in_paragraph(para_elem, eng_text):
    """
    Check if the English translation text already exists in the paragraph.
    This prevents duplicates when the original cell already has bilingual
    content (EN + CH in the same paragraph, separated by line breaks).
    """
    if not eng_text:
        return False
    eng_lower = eng_text.strip().lower()
    # Collect all text from this paragraph's runs
    existing_parts = []
    for t in para_elem.findall(f".//{qn('w:t')}"):
        if t.text:
            existing_parts.append(t.text)
    existing_full = "".join(existing_parts).strip().lower()
    # Check if translation already appears in the paragraph text
    return eng_lower in existing_full


def paragraph_is_already_bilingual(para_elem):
    """
    Check if a paragraph already contains both English and Chinese text
    (e.g., 'Doc. Title\\n文件标题' in the same paragraph via w:br).
    If so, adding another English line would be a duplicate.
    """
    all_text = []
    for t in para_elem.findall(f".//{qn('w:t')}"):
        if t.text:
            all_text.append(t.text)
    full = "".join(all_text)
    return has_chinese(full) and has_english(full)


def make_english_run(eng_text, is_first=False):
    elements = []
    if not is_first:
        br = etree.Element(qn("w:br"))
        br_run = etree.Element(qn("w:r"))
        br_run.append(br)
        elements.append(br_run)
    r = etree.Element(qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), "21")
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), "21")
    i_elem = etree.SubElement(rPr, qn("w:i"))
    i_elem.set(qn("w:val"), "false")
    iCs = etree.SubElement(rPr, qn("w:iCs"))
    iCs.set(qn("w:val"), "false")
    t = etree.SubElement(r, qn("w:t"))
    t.text = eng_text
    t.set(qn("xml:space"), "preserve")
    elements.append(r)
    return elements


def get_element_path(elem, tree):
    """Get a unique XPath for an element within the document tree."""
    try:
        return tree.getpath(elem)
    except Exception:
        return None


def insert_translations_safe(input_path, translations_path, output_path):
    doc = docx.Document(input_path)
    tree = doc.element.getroottree()

    with open(translations_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    para_translations = {}
    for item in data.get("paragraphs", []):
        idx = item.get("index", -1)
        eng = item.get("translation", "").strip()
        if idx >= 0 and eng:
            para_translations[idx] = eng

    table_translations = {}
    for table_item in data.get("tables", []):
        ti = table_item.get("index", -1)
        if ti < 0:
            continue
        table_translations[ti] = []
        for row_item in table_item.get("rows", []):
            ri = row_item.get("index", -1)
            for cell_item in row_item.get("cells", []):
                eng = cell_item.get("translation", "").strip()
                if eng:
                    table_translations[ti].append((
                        ri, cell_item["cell_index"], cell_item["para_index"], eng
                    ))

    # --- Body paragraphs ---
    body = doc.element.body
    paragraph_elements = body.findall(qn("w:p"))
    inserted = 0
    para_skipped = 0
    for i, p_elem in enumerate(paragraph_elements):
        if i in para_translations:
            eng_text = para_translations[i]
            # Skip if paragraph element already contains mixed EN+CH (already bilingual)
            all_text = "".join(
                t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")
            ).strip()
            if is_already_bilingual_v2(all_text):
                para_skipped += 1
                continue
            # Skip if translation text already appears in paragraph
            if translation_already_in_paragraph(p_elem, eng_text):
                para_skipped += 1
                continue
            for elem in make_english_run(eng_text, is_first=(not all_text)):
                p_elem.append(elem)
            inserted += 1

    # --- Table cells with XPath-based merged cell deduplication ---
    # Also skip cells where English already exists above Chinese (bilingual headers)
    modified_paths = set()
    table_inserted = 0
    skipped = 0

    for ti, translations_list in table_translations.items():
        if ti >= len(doc.tables):
            continue
        table = doc.tables[ti]
        for ri, ci, pi, eng_text in translations_list:
            if ri >= len(table.rows):
                continue
            row = table.rows[ri]
            if ci >= len(row.cells):
                continue
            cell = row.cells[ci]
            if pi < len(cell.paragraphs):
                target_p = cell.paragraphs[pi]
                target_elem = target_p._element
                # Get full text of target paragraph for already-bilingual check
                target_text = "".join(
                    t.text or "" for t in target_elem.findall(f".//{qn('w:t')}")
                ).strip()
                # Skip if paragraph already has mixed EN words + CH (already bilingual)
                if is_already_bilingual_v2(target_text):
                    skipped += 1
                    continue
                # Skip if this cell already has English text above the Chinese
                # (original document is already bilingual in this cell)
                if cell_has_existing_english_above(cell, pi):
                    skipped += 1
                    continue
                # Skip if translation already appears in this paragraph
                # (original already has bilingual EN+CH in same paragraph)
                if translation_already_in_paragraph(target_elem, eng_text):
                    skipped += 1
                    continue
                elem_path = get_element_path(target_elem, tree)
                if elem_path and elem_path in modified_paths:
                    continue
                existing_text = "".join(
                    t.text or "" for t in target_elem.findall(f".//{qn('w:t')}")
                ).strip()
                for elem in make_english_run(eng_text, is_first=(not existing_text)):
                    target_elem.append(elem)
                if elem_path:
                    modified_paths.add(elem_path)
                table_inserted += 1

    out = output_path or input_path
    doc.save(out)
    print(f"Inserted {inserted} paragraph translations (skipped {para_skipped}), "
          f"{table_inserted} table cell translations "
          f"(skipped {skipped} already-bilingual cells) into {out}")
    return out


if __name__ == "__main__":
    inp = sys.argv[1]
    trans = sys.argv[2]
    out = None
    args = sys.argv[3:]
    if "--output" in args:
        idx = args.index("--output")
        if idx + 1 < len(args):
            out = args[idx + 1]
    insert_translations_safe(inp, trans, out)

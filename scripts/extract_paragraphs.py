"""
Extract all translatable content from a .docx file as JSON.
Run this first to get the content for translation.
Usage: python extract_paragraphs.py <input.docx> [--output text.json]
"""
import json
import sys
import docx


def extract_paragraphs(input_path, output_path=None):
    doc = docx.Document(input_path)
    paragraphs = []
    tables = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "style": p.style.name if p.style else "",
                "text": text
            })

    for ti, table in enumerate(doc.tables):
        table_data = {"index": ti, "rows": []}
        for ri, row in enumerate(table.rows):
            row_data = {"index": ri, "cells": []}
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    text = p.text.strip()
                    if text:
                        row_data["cells"].append({
                            "cell_index": ci,
                            "para_index": pi,
                            "text": text
                        })
            if row_data["cells"]:
                table_data["rows"].append(row_data)
        if table_data["rows"]:
            tables.append(table_data)

    result = {"paragraphs": paragraphs, "tables": tables}
    output = output_path or input_path.replace(".docx", "_extracted.json")
    with open(output, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Extracted {len(paragraphs)} paragraphs, {len(tables)} tables → {output}")
    return output


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_paragraphs.py <input.docx> [--output text.json]")
        sys.exit(1)
    input_file = sys.argv[1]
    out = None
    if len(sys.argv) > 2 and sys.argv[2] == "--output" and len(sys.argv) > 3:
        out = sys.argv[3]
    extract_paragraphs(input_file, out)

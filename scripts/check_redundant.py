import docx

doc = docx.Document(r'C:\Users\Adam Cheng\WorkBuddy\Chinese to English\101-SMP-000020_1.0_Test 1_EN.docx')

print('=== Checking all table cells for redundant English translations ===')
redundant = []

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            paras = cell.paragraphs
            # Check all paragraphs in this cell
            for pi, p in enumerate(paras):
                text = p.text.strip()
                if not text:
                    continue
                lines = text.split('\n')
                # Pattern: line0=English, line1=Chinese, line2=English(dup)
                # This happens when insert added English after Chinese that
                # already had English above it
                if len(lines) >= 3:
                    l0_chn = any(ord(c) > 0x4e00 for c in lines[0])
                    l1_chn = any(ord(c) > 0x4e00 for c in lines[1])
                    l2_chn = any(ord(c) > 0x4e00 for c in lines[2])
                    if not l0_chn and l1_chn and not l2_chn:
                        redundant.append((ti, ri, ci, pi, lines[0][:40], lines[1][:40], lines[2][:40]))

# Also check: cell has English para above, Chinese para, then English added below
print('=== Broader check: cells with 2+ English lines and 1+ Chinese ===')
broader = []
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            paras = cell.paragraphs
            if len(paras) < 2:
                continue
            eng_count = 0
            chn_count = 0
            eng_texts = []
            chn_texts = []
            for pi, p in enumerate(paras):
                text = p.text.strip()
                if not text:
                    continue
                lines = text.split('\n')
                for line in lines:
                    if not line.strip():
                        continue
                    has_chn = any(ord(c) > 0x4e00 for c in line)
                    if has_chn:
                        chn_count += 1
                        chn_texts.append((pi, line[:50]))
                    else:
                        eng_count += 1
                        eng_texts.append((pi, line[:50]))
            if eng_count >= 2 and chn_count >= 1:
                broader.append((ti, ri, ci, eng_texts, chn_texts))

if redundant:
    print(f'\nFound {len(redundant)} redundant translations (EN-CHN-EN pattern in same paragraph):')
    for ti, ri, ci, pi, l0, l1, l2 in redundant:
        print(f'  T{ti} R{ri} C{ci} P{pi}:')
        print(f'    Line0 (EN orig):  {l0}')
        print(f'    Line1 (CHN orig): {l1}')
        print(f'    Line2 (EN dup):   {l2}')
else:
    print('No EN-CHN-EN pattern in single paragraphs.')

if broader:
    print(f'\nFound {len(broader)} cells with 2+ English lines:')
    for ti, ri, ci, eng_texts, chn_texts in broader:
        print(f'  T{ti} R{ri} C{ci}:')
        for pi, t in eng_texts:
            print(f'    EN P{pi}: {t}')
        for pi, t in chn_texts:
            print(f'    CHN P{pi}: {t}')
else:
    print('No cells with 2+ English lines found.')

"""
Replace existing English in a bilingual .docx with corrected English.

Consumes corrections.json produced by correct_english.py. Unlike
insert_translations_safe.py (which ADDS English below Chinese), this script
REPLACES the existing English text in place, keeping the Chinese untouched.

Two replacement modes (per segment, from corrections.json):
  mode "same" — the paragraph / cell paragraph contains both Chinese and
                English (typically "CN\nEN"); only the English portion is
                replaced.
  mode "next" — the English lives in the paragraph right below the Chinese
                paragraph (or the next paragraph inside a table cell); that
                whole paragraph is replaced with the corrected English.

Merged table cells are deduplicated via element paths so each cell is only
modified once.

Usage:
    python insert_corrections_safe.py <input.docx> <corrections.json> [--output out.docx]
"""
import copy
import datetime
import json
import sys
from lxml import etree
import docx
from docx.oxml.ns import qn


# ── Track changes helpers ─────────────────────────────────────────────

_change_id_counter = [0]


def _next_change_id():
    _change_id_counter[0] += 1
    return str(_change_id_counter[0])


def _get_timestamp():
    return datetime.datetime.now(datetime.timezone.utc).strftime(
        "%Y-%m-%dT%H:%M:%SZ"
    )


def _convert_run_to_deleted(run_elem):
    """Deep-copy a run and convert its w:t to w:delText (for w:del)."""
    r = copy.deepcopy(run_elem)
    for child in r:
        if child.tag == qn("w:t"):
            child.tag = qn("w:delText")
    return r


def _wrap_runs_in_del(runs, author="DocTranslator"):
    """Wrap deep-copied runs (w:t→w:delText) in a single w:del element."""
    if not runs:
        return None
    del_elem = etree.Element(qn("w:del"))
    del_elem.set(qn("w:id"), _next_change_id())
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), _get_timestamp())
    for r in runs:
        del_elem.append(_convert_run_to_deleted(r))
    return del_elem


def _wrap_runs_in_ins(runs, author="DocTranslator"):
    """Wrap runs in a single w:ins element."""
    if not runs:
        return None
    ins_elem = etree.Element(qn("w:ins"))
    ins_elem.set(qn("w:id"), _next_change_id())
    ins_elem.set(qn("w:author"), author)
    ins_elem.set(qn("w:date"), _get_timestamp())
    for r in runs:
        ins_elem.append(r)
    return ins_elem


def _make_deleted_partial_runs(run_elem, text_start, text_end):
    """Create run(s) containing text[text_start:text_end] of *run_elem*,
    using w:delText (for use inside w:del). Preserves line breaks as w:br.

    Returns a list of w:r elements (empty list if the slice is empty).
    """
    full_text = _run_text(run_elem)
    del_text = full_text[text_start:text_end]
    if not del_text:
        return []

    rpr = _get_rpr_from_run(run_elem)
    runs = []
    lines = del_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br_run = etree.Element(qn("w:r"))
            if rpr is not None:
                br_run.append(copy.deepcopy(rpr))
            br_run.append(etree.Element(qn("w:br")))
            runs.append(br_run)
        if line:
            r = etree.Element(qn("w:r"))
            if rpr is not None:
                r.append(copy.deepcopy(rpr))
            dt = etree.SubElement(r, qn("w:delText"))
            dt.text = line
            dt.set(qn("xml:space"), "preserve")
            runs.append(r)
    return runs


def has_chinese(text):
    if not text:
        return False
    return any(ord(c) > 0x4e00 for c in text)


def has_english(text):
    if not text:
        return False
    return any(c.isalpha() and ord(c) < 128 for c in text)


# ── Low-level XML helpers ──────────────────────────────────────────────

def _run_text(run_elem):
    """Rendered text of a run (w:br -> '\\n', w:tab -> '\\t')."""
    parts = []
    for child in run_elem:
        tag = child.tag
        if tag == qn("w:t"):
            parts.append(child.text or "")
        elif tag == qn("w:br"):
            parts.append("\n")
        elif tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


def _get_rpr_from_run(run_elem):
    """Return the w:rPr element from a run, or None if it has none."""
    for child in run_elem:
        if child.tag == qn("w:rPr"):
            return child
    return None


def _text_run(text, rpr_template=None):
    """An English run.

    If *rpr_template* (a w:rPr element from the original English run) is
    provided, it is deep-copied so the corrected text inherits the exact
    same formatting (font family, size, color, etc.) as the original
    English it replaces.

    Otherwise a Times New Roman 10.5pt (21 half-points), non-italic run is
    created as a fallback.
    """
    r = etree.Element(qn("w:r"))
    if rpr_template is not None:
        r.append(copy.deepcopy(rpr_template))
    else:
        rPr = etree.SubElement(r, qn("w:rPr"))
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        sz = etree.SubElement(rPr, qn("w:sz"))
        sz.set(qn("w:val"), "21")
        szCs = etree.SubElement(rPr, qn("w:szCs"))
        szCs.set(qn("w:val"), "21")
        i_elem = etree.SubElement(rPr, qn("w:i"))
        i_elem.set(qn("w:val"), "false")
        iCs = etree.SubElement(rPr, qn("w:iCs"))
        iCs.set(qn("w:val"), "false")
    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")
    return r


def make_corrected_runs(eng_text, rpr_template=None):
    """Runs for corrected English, inserting w:br between lines.

    *rpr_template* (if provided) is passed through to _text_run so every
    text run inherits the original English formatting (font size, family,
    etc.).
    """
    lines = str(eng_text).split("\n")
    runs = []
    for i, line in enumerate(lines):
        if i > 0:
            br_run = etree.Element(qn("w:r"))
            br_run.append(etree.Element(qn("w:br")))
            runs.append(br_run)
        if line:
            runs.append(_text_run(line, rpr_template))
    if not runs:
        runs.append(_text_run("", rpr_template))
    return runs


def _paragraph_text(p_elem):
    return "".join(
        t.text or "" for t in p_elem.findall(f".//{qn('w:t')}")
    ).strip()


def _locate_en_region(full_text, cn_part, en_part):
    """Find (start, end) offsets of the English portion within a paragraph's
    full rendered text. Handles both 'CN below EN' and 'EN below CN' layouts.
    Returns None if the Chinese text cannot be located."""
    if not full_text:
        return None
    cn = (cn_part or "").strip()
    en = (en_part or "").strip()
    cn_pos = full_text.find(cn)
    if cn_pos < 0:
        return None

    if en:
        en_pos = full_text.find(en)
        if en_pos >= 0:
            if en_pos < cn_pos:
                # English above Chinese: replace [en_pos, cn_pos)
                start, end = en_pos, cn_pos
                while start < end and full_text[start] in " \t\n":
                    start += 1
                # keep any separator (newline/space) between the English and
                # the Chinese so the corrected English stays on its own line
                while end > start and full_text[end - 1] in " \t\n":
                    end -= 1
                return (start, end)
            # English below Chinese: replace [en_pos, end of paragraph)
            start = en_pos
            return (start, len(full_text))

    # No English marker found: replace everything after the Chinese
    start = cn_pos + len(cn)
    while start < len(full_text) and full_text[start] in " \t\n":
        start += 1
    return (start, len(full_text))


def _truncate_run_at(run_elem, keep_len):
    """Keep only the first keep_len characters of a run's rendered text."""
    if keep_len <= 0:
        for item in [c for c in run_elem if c.tag in (qn("w:t"), qn("w:br"), qn("w:tab"))]:
            run_elem.remove(item)
        return
    items = [c for c in run_elem if c.tag in (qn("w:t"), qn("w:br"), qn("w:tab"))]
    offset = 0
    for item in items:
        if item.tag == qn("w:t"):
            item_len = len(item.text or "")
        else:
            item_len = 1
        istart, iend = offset, offset + item_len
        offset = iend
        if iend <= keep_len:
            continue
        if istart >= keep_len:
            run_elem.remove(item)
        else:
            if item.tag == qn("w:t"):
                keep_in_item = keep_len - istart
                item.text = (item.text or "")[:keep_in_item]
            else:
                run_elem.remove(item)


def _keep_run_suffix(run_elem, keep_from):
    """Keep only the rendered text at offsets >= keep_from of a run."""
    items = [c for c in run_elem if c.tag in (qn("w:t"), qn("w:br"), qn("w:tab"))]
    offset = 0
    for item in items:
        if item.tag == qn("w:t"):
            item_len = len(item.text or "")
        else:
            item_len = 1
        istart, iend = offset, offset + item_len
        offset = iend
        if iend <= keep_from:
            run_elem.remove(item)
        elif istart < keep_from:
            if item.tag == qn("w:t"):
                item.text = (item.text or "")[keep_from - istart:]
            else:
                run_elem.remove(item)
        # items fully after keep_from stay untouched


def replace_english_in_paragraph(p_elem, cn_part, en_part, corrected_en,
                                  track_changes=False):
    """Replace the English portion of a mixed CN+EN paragraph with the
    corrected English, preserving the Chinese runs (before AND after the
    English) together with their formatting.

    The run properties (font size, family, etc.) of the first English-
    region run are captured and applied to the corrected text so it
    matches the original English's appearance.

    When *track_changes* is True, the old English text is wrapped in
    w:del and the new text in w:ins so Word displays tracked changes.
    """
    runs = [r for r in p_elem if r.tag == qn("w:r")]
    full = "".join(_run_text(r) for r in runs)
    region = _locate_en_region(full, cn_part, en_part)
    if region is None:
        return False
    start, end = region

    offset = 0
    to_remove = []
    prefix_run = None    # run straddling the region start, truncated to its prefix
    suffix_run = None    # run straddling the region end, truncated to its suffix
    suffix_clone = None  # deepcopy of a single run spanning the whole region (suffix part)
    first_tail_run = None  # first run fully after the region (insert point)
    rpr_template = None  # rPr captured from the first English-region run
    deleted_runs = []    # runs to wrap in w:del (for track changes)

    for r in runs:
        rtext = _run_text(r)
        rlen = len(rtext)
        rstart, rend = offset, offset + rlen
        offset = rend
        if rstart >= end:
            if first_tail_run is None:
                first_tail_run = r
            continue                       # fully after the region — keep
        if rend <= start:
            continue                       # fully before the region — keep
        # The run intersects [start, end) — this is an English-region run.
        # Capture its rPr (first one wins) so the corrected text inherits
        # the original English font size and formatting.
        if rpr_template is None:
            rpr_template = _get_rpr_from_run(r)
        if rstart < start:
            if rend > end:
                # a single run spans the whole region — clone for the suffix
                clone = copy.deepcopy(r)
                _keep_run_suffix(clone, end - rstart)
                if _run_text(clone):
                    suffix_clone = clone
                # Capture deleted portion [start, end) for track changes
                if track_changes:
                    deleted_runs.extend(
                        _make_deleted_partial_runs(r, start - rstart, end - rstart)
                    )
                _truncate_run_at(r, start - rstart)
                prefix_run = r
                if not _run_text(r):
                    to_remove.append(r)
                    prefix_run = None
            else:
                # straddles start only — deleted portion [start, rend)
                if track_changes:
                    deleted_runs.extend(
                        _make_deleted_partial_runs(r, start - rstart, rlen)
                    )
                _truncate_run_at(r, start - rstart)
                prefix_run = r
                if not _run_text(r):
                    to_remove.append(r)
                    prefix_run = None
        elif rend > end:
            # straddles end only — deleted portion [rstart, end) = [0, end-rstart)
            if track_changes:
                deleted_runs.extend(
                    _make_deleted_partial_runs(r, 0, end - rstart)
                )
            _keep_run_suffix(r, end - rstart)
            suffix_run = r
            if not _run_text(r):
                to_remove.append(r)
                suffix_run = None
        else:
            # fully inside the region — remove
            if track_changes:
                deleted_runs.append(_convert_run_to_deleted(r))
            to_remove.append(r)

    for r in to_remove:
        p_elem.remove(r)

    new_runs = make_corrected_runs(corrected_en, rpr_template)

    # Build w:del and w:ins elements for track changes
    del_elem = None
    ins_elem = None
    if track_changes:
        del_elem = _wrap_runs_in_del(deleted_runs) if deleted_runs else None
        ins_elem = _wrap_runs_in_ins(new_runs)
    else:
        ins_elem = None

    # Determine insertion anchor and order
    # For track changes: insert [del_elem, ins_elem] at the anchor
    # For non-track-changes: insert new_runs at the anchor (original behavior)
    def _insert_elements(elements):
        """Insert a list of elements at the correct position in the paragraph."""
        if suffix_run is not None:
            for el in elements:
                suffix_run.addprevious(el)
        elif suffix_clone is not None:
            anchor = prefix_run
            if anchor is not None:
                for el in elements:
                    anchor.addnext(el)
                    anchor = el
                anchor.addnext(suffix_clone)
            else:
                for el in elements:
                    p_elem.append(el)
                p_elem.append(suffix_clone)
        elif first_tail_run is not None:
            for el in elements:
                first_tail_run.addprevious(el)
        else:
            for el in elements:
                p_elem.append(el)

    if track_changes:
        elements = []
        if del_elem is not None:
            elements.append(del_elem)
        if ins_elem is not None:
            elements.append(ins_elem)
        _insert_elements(elements)
    else:
        _insert_elements(new_runs)
    return True


def set_paragraph_text(p_elem, text, track_changes=False):
    """Replace all runs of a paragraph with the corrected English text.

    The formatting (font size, family, etc.) of the first existing run is
    captured *before* the runs are removed, so the corrected text inherits
    the same appearance as the original English it replaces.

    When *track_changes* is True, old runs are wrapped in w:del and new
    runs in w:ins so Word displays them as tracked changes.
    """
    runs = [r for r in p_elem if r.tag == qn("w:r")]
    # Capture rPr from the first run that has one, so the corrected text
    # keeps the original English font size and other formatting.
    rpr_template = None
    for r in runs:
        rpr = _get_rpr_from_run(r)
        if rpr is not None:
            rpr_template = rpr
            break

    if track_changes and runs:
        # Wrap old runs in w:del (deep-copied, w:t→w:delText)
        del_elem = _wrap_runs_in_del(runs)
        # Remove old runs
        for r in runs:
            p_elem.remove(r)
        # Add w:del then w:ins
        if del_elem is not None:
            p_elem.append(del_elem)
        new_runs = make_corrected_runs(text, rpr_template)
        ins_elem = _wrap_runs_in_ins(new_runs)
        if ins_elem is not None:
            p_elem.append(ins_elem)
    else:
        for r in runs:
            p_elem.remove(r)
        for run in make_corrected_runs(text, rpr_template):
            p_elem.append(run)


def get_element_path(elem, tree):
    try:
        return tree.getpath(elem)
    except Exception:
        return None


# ── Main ───────────────────────────────────────────────────────────────

def insert_corrections_safe(input_path, corrections_path, output_path=None,
                             track_changes=False):
    doc = docx.Document(input_path)
    tree = doc.element.getroottree()

    with open(corrections_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # --- Body paragraphs ---
    # Use doc.paragraphs (body-level only) to match the indexing used by
    # extract_paragraphs.py — body.findall(qn("w:p")) would also return the
    # paragraphs nested inside tables and shift every index.
    paragraph_elements = [p._element for p in doc.paragraphs]
    replaced = 0
    para_skipped = 0

    for item in data.get("paragraphs", []):
        idx = item.get("replace_index", -1)
        mode = item.get("replace_mode", "next")
        eng = (item.get("translation") or "").strip()
        if idx < 0 or idx >= len(paragraph_elements) or not eng:
            continue
        if eng.startswith("[CORRECTION ERROR"):
            para_skipped += 1
            continue
        p_elem = paragraph_elements[idx]
        if mode == "same":
            ok = replace_english_in_paragraph(
                p_elem, item.get("text", ""), item.get("original_en", ""), eng,
                track_changes=track_changes,
            )
            if not ok:
                para_skipped += 1
                continue
        else:
            set_paragraph_text(p_elem, eng, track_changes=track_changes)
        replaced += 1

    # --- Table cells (with merged-cell dedup) ---
    modified_paths = set()
    cell_replaced = 0
    cell_skipped = 0

    for table_item in data.get("tables", []):
        ti = table_item.get("index", -1)
        if ti < 0 or ti >= len(doc.tables):
            continue
        table = doc.tables[ti]
        for row_item in table_item.get("rows", []):
            ri = row_item.get("index", -1)
            if ri < 0 or ri >= len(table.rows):
                continue
            row = table.rows[ri]
            for cell_item in row_item.get("cells", []):
                eng = (cell_item.get("translation") or "").strip()
                if not eng or eng.startswith("[CORRECTION ERROR"):
                    cell_skipped += 1
                    continue
                ci = cell_item.get("cell_index", -1)
                replace_pi = cell_item.get("replace_pi", -1)
                mode = cell_item.get("replace_mode", "next")
                if ci < 0 or ci >= len(row.cells):
                    continue
                cell = row.cells[ci]
                if replace_pi < 0 or replace_pi >= len(cell.paragraphs):
                    continue
                target_p = cell.paragraphs[replace_pi]
                target_elem = target_p._element
                elem_path = get_element_path(target_elem, tree)
                if elem_path and elem_path in modified_paths:
                    continue
                if mode == "same":
                    ok = replace_english_in_paragraph(
                        target_elem,
                        cell_item.get("text", ""),
                        cell_item.get("original_en", ""),
                        eng,
                        track_changes=track_changes,
                    )
                    if not ok:
                        cell_skipped += 1
                        continue
                else:
                    set_paragraph_text(target_elem, eng, track_changes=track_changes)
                if elem_path:
                    modified_paths.add(elem_path)
                cell_replaced += 1

    out = output_path or input_path
    doc.save(out)
    mode_label = " (track changes)" if track_changes else ""
    print(
        f"Replaced {replaced} paragraph corrections (skipped {para_skipped}), "
        f"{cell_replaced} table cell corrections (skipped {cell_skipped}){mode_label} → {out}"
    )
    return out


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python insert_corrections_safe.py <input.docx> <corrections.json> [--output out.docx] [--track-changes]")
        sys.exit(1)
    inp = sys.argv[1]
    corr = sys.argv[2]
    out = None
    tc = False
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            out = sys.argv[idx + 1]
    if "--track-changes" in sys.argv:
        tc = True
    insert_corrections_safe(inp, corr, out, track_changes=tc)
